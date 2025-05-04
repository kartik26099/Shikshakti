import os
from typing import List, BinaryIO
import PyPDF2
import docx
import re
from sentence_transformers import SentenceTransformer
from db_setup import SessionLocal, Document, DocumentChunk
from db_setup import Base, engine
from functools import lru_cache

# Initialize database tables
Base.metadata.create_all(engine)

# Load embedding model (using sentence-transformers instead of OpenAI)
model = SentenceTransformer('all-MiniLM-L6-v2')
@lru_cache(maxsize=1000)
def extract_text_from_file(file: BinaryIO, filename: str) -> str:
    """Extract text from uploaded files based on file type."""
    if filename.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    elif filename.endswith('.docx'):
        doc = docx.Document(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    elif filename.endswith('.txt'):
        return file.read().decode('utf-8')
    else:
        raise ValueError(f"Unsupported file format: {filename}")

def split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 150) -> List[str]:
    """Split text into overlapping chunks with a smaller size."""
    print(f"Starting split_text_into_chunks with text of length {len(text)}")
    
    if not text:
        print("Text is empty, returning empty list")
        return []
    
    # For very small texts, just return the whole text as one chunk
    if len(text) <= chunk_size:
        print(f"Text length ({len(text)}) is smaller than chunk_size ({chunk_size}), returning whole text as one chunk")
        return [text]
    
    # Clean text
    print("Cleaning text...")
    text = re.sub(r'\s+', ' ', text).strip()
    print(f"Cleaned text length: {len(text)}")
    
    # Split into chunks
    chunks = []
    start = 0
    print(f"Starting chunking loop with start={start}, len(text)={len(text)}")
    
    while start < len(text):
        end = min(start + chunk_size, len(text))
        print(f"Iteration: start={start}, end={end}")
        
        # Find a good breaking point (space) if not at the end
        if end < len(text):
            print(f"Finding breaking point near position {end}")
            original_end = end
            while end > start and text[end] != ' ':
                end -= 1
                
            if end == start:  # No space found, just cut at chunk_size
                print(f"No space found between {start} and {original_end}, using original end")
                end = original_end
            else:
                print(f"Found space at position {end}")
        
        chunk = text[start:end]
        print(f"Adding chunk of length {len(chunk)}")
        chunks.append(chunk)
        
        new_start = end - overlap
        print(f"New start position: {new_start} (end={end}, overlap={overlap})")
        
        # Prevent infinite loop
        if new_start <= start:
            print(f"WARNING: New start position ({new_start}) <= current start ({start}), advancing to end")
            new_start = end
            
        start = new_start
        
        # Break if we've reached the end
        if end >= len(text):
            print("Reached end of text, breaking loop")
            break
    
    print(f"Finished chunking, created {len(chunks)} chunks")
    return chunks


def generate_embedding(text: str) -> List[float]:
    """Generate embedding vector for text using sentence-transformers with caching."""
    try:
        return model.encode(text).tolist()
    except Exception as e:
        print(f"Error generating embedding: {str(e)}")
        # Return a zero vector as fallback
        return [0.0] * 384  # Default dimension for 'all-MiniLM-L6-v2'

def process_document(file: BinaryIO, filename: str, title: str) -> int:
    """Process document, extract text, and store chunks without using Pinecone."""
    try:
        print(f"Starting to process document: {filename}")
        print(f"Document title: {title}")
        
        # Extract text from file
        print("Extracting text from file...")
        text = extract_text_from_file(file, filename)
        print(f"Extracted text length: {len(text)} characters")
        
        # Create document in database
        print("Creating document in database...")
        db = SessionLocal()
        try:
            document = Document(title=title, content=text)
            db.add(document)
            db.commit()
            db.refresh(document)
            print(f"Document created with ID: {document.id}")
            
            # Split text into chunks
            print(f"Splitting text into chunks (size={len(text)}, chunk_size=1000, overlap=150)...")
            chunks = split_text_into_chunks(text, chunk_size=1000, overlap=150)
            print(f"Created {len(chunks)} chunks")
            
            # Store full chunks in database
            print("Storing chunks in database...")
            store_full_chunks(document.id, chunks)
            print("Chunks stored successfully")
            
            return document.id
        finally:
            db.close()
    except Exception as e:
        print(f"ERROR in process_document: {str(e)}")
        import traceback
        traceback.print_exc()
        raise  # Re-raise the exception to be caught by the endpoint handler

def store_full_chunks(document_id: int, chunks: List[str]) -> None:
    """Store full text chunks in database for retrieval."""
    print(f"Storing {len(chunks)} chunks for document ID {document_id}")
    
    db = SessionLocal()
    try:
        # Delete any existing chunks for this document
        print(f"Deleting existing chunks for document ID {document_id}")
        db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
        
        # Insert new chunks
        print("Inserting new chunks...")
        for i, chunk in enumerate(chunks):
            print(f"Creating chunk record {i+1}/{len(chunks)} with length {len(chunk)}")
            chunk_record = DocumentChunk(
                document_id=document_id,
                chunk_index=i,
                content=chunk
            )
            db.add(chunk_record)
        
        print("Committing chunks to database...")
        db.commit()
        print("Chunks committed successfully")
    except Exception as e:
        db.rollback()
        print(f"ERROR storing document chunks: {str(e)}")
        import traceback
        traceback.print_exc()
        raise
    finally:
        db.close()

def cosine_similarity(vec1, vec2):
    """Calculate cosine similarity between two vectors."""
    dot_product = sum(a*b for a, b in zip(vec1, vec2))
    magnitude1 = sum(a*a for a in vec1) ** 0.5
    magnitude2 = sum(b*b for b in vec2) ** 0.5
    if magnitude1 * magnitude2 == 0:
        return 0
    return dot_product / (magnitude1 * magnitude2)